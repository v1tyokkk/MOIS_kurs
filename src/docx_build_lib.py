#!/usr/bin/env python3
"""Сборка DOCX из *.txt курсовой (KURS / KURS2) через word-manager."""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

WORD_MANAGER_SRC = Path(__file__).resolve().parent.parent.parent / "word manager" / "src"
if not WORD_MANAGER_SRC.is_dir():
    WORD_MANAGER_SRC = Path("/Users/roman/Documents/word manager/src")
if str(WORD_MANAGER_SRC) not in sys.path:
    sys.path.insert(0, str(WORD_MANAGER_SRC))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

from word_manager import (
    WordBuildState,
    WordDocumentConfig,
    add_heading,
    add_paragraph,
    add_rectangle_border,
    create_document,
    save_docx,
)

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 13
FONT_COLOR = RGBColor(0, 0, 0)
LINE_SPACING = 1.2
FIRST_LINE_INDENT_CM = 2
TITLE_SPACE_PT = 39
TOC_START = "СОДЕРЖАНИЕ"

# Раздел — CAPS целиком или «1 СИСТЕМНЫЙ АНАЛИЗ…» (KURS)
# Глава — x.y. не капсом (1.1., 2.3.)
RE_CHAPTER = re.compile(r"^\d+\.\d+\.\s+")
RE_NUMBERED_SECTION = re.compile(r"^\d+\s+(.+)$")
# Подписи осей / шапка CSV в коде: X,Y,Z или "X", "Y", "Z"
RE_AXIS_CSV_LABELS = re.compile(
    r'^["\']?[XYZ]["\']?(?:\s*,\s*["\']?[XYZ]["\']?){0,2}\s*,?\s*$',
    re.IGNORECASE,
)

MARGIN_LEFT_MM = 25
MARGIN_TOP_MM = 15
MARGIN_RIGHT_MM = 10
MARGIN_BOTTOM_MM = 35

BORDER_LEFT_MM = 20
BORDER_TOP_MM = 5
BORDER_RIGHT_MM = 5
BORDER_BOTTOM_MM = 5

TXT_FILES: list[tuple[str, str]] = [
    ("ПЗ.txt", "ПЗ"),
    ("ПЗ_А.txt", "ПЗ_А"),
    ("ПЗ_Б.txt", "ПЗ_Б"),
    ("ПЗ_В.txt", "ПЗ_В"),
]

COMBINED_STEM = "ПЗ_полная"
COMBINED_TXT_ORDER = [name for name, _ in TXT_FILES]


def is_heading_line(line: str) -> bool:
    """Строки с «=», подписи X,Y,Z — не раздел и не глава."""
    s = line.strip()
    if "=" in s:
        return False
    if RE_AXIS_CSV_LABELS.match(s):
        return False
    return True


def is_all_caps_title(line: str) -> bool:
    s = line.strip()
    if not is_heading_line(s):
        return False
    if len(s) < 2:
        return False
    letters = [ch for ch in s if ch.isalpha()]
    if not letters:
        return False
    return all(ch.isupper() for ch in letters)


def is_section_caps_title(line: str) -> bool:
    """Раздел: ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, 1 СИСТЕМНЫЙ АНАЛИЗ… (KURS)."""
    s = line.strip()
    if not is_heading_line(s):
        return False
    if is_all_caps_title(s):
        return True
    m = RE_NUMBERED_SECTION.match(s)
    if m:
        return is_all_caps_title(m.group(1))
    return False


def is_chapter_heading(line: str) -> bool:
    s = line.strip()
    if not is_heading_line(s):
        return False
    return bool(RE_CHAPTER.match(s))


def apply_page_setup(doc: Document) -> None:
    sec = doc.sections[0]
    sec.left_margin = Mm(MARGIN_LEFT_MM)
    sec.top_margin = Mm(MARGIN_TOP_MM)
    sec.right_margin = Mm(MARGIN_RIGHT_MM)
    sec.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    sec.header_distance = Mm(0)
    sec.footer_distance = Mm(0)


def apply_page_border(doc: Document, state: WordBuildState | None = None) -> None:
    add_rectangle_border(
        doc,
        left_border_from_left=BORDER_LEFT_MM,
        top_border_from_top=BORDER_TOP_MM,
        right_border_from_right=BORDER_RIGHT_MM,
        bottom_border_from_bottom=BORDER_BOTTOM_MM,
        state=state,
    )


def _set_rfonts(r_pr, name: str) -> None:
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(key), name)


def apply_base_styles(doc: Document) -> None:
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = FONT_NAME
        st.font.size = Pt(FONT_SIZE_PT)
        st.font.color.rgb = FONT_COLOR
        st.font.bold = style_name in ("Heading 1", "Heading 2")
        _set_rfonts(st.element.get_or_add_rPr(), FONT_NAME)
        pf = st.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = LINE_SPACING
        if style_name in ("Normal", "Heading 1", "Heading 2"):
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if style_name == "Normal":
            pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
        elif style_name == "Heading 1":
            pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
            pf.space_after = Pt(TITLE_SPACE_PT)
        elif style_name == "Heading 2":
            pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
            pf.space_before = Pt(TITLE_SPACE_PT)
            pf.space_after = Pt(TITLE_SPACE_PT)


def _is_heading_paragraph(para) -> bool:
    name = para.style.name if para.style else ""
    return name in ("Heading 1", "Heading 2")


def format_all_runs(doc: Document) -> None:
    for para in doc.paragraphs:
        if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bold = _is_heading_paragraph(para)
        for run in para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_PT)
            run.font.color.rgb = FONT_COLOR
            run.font.bold = bold
            _set_rfonts(run._element.get_or_add_rPr(), FONT_NAME)


def apply_title_spacing(doc: Document) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        pf = para.paragraph_format
        if style == "Heading 2" and is_chapter_heading(text):
            pf.space_before = Pt(TITLE_SPACE_PT)
            pf.space_after = Pt(TITLE_SPACE_PT)
        elif style == "Heading 1" and is_section_caps_title(text):
            pf.space_before = Pt(0)
            pf.space_after = Pt(TITLE_SPACE_PT)


def apply_paragraph_format(
    para,
    *,
    body: bool = True,
    page_break_before: bool = False,
    section_caps: bool = False,
    chapter_heading: bool = False,
    toc_entry: bool = False,
) -> None:
    pf = para.paragraph_format
    if section_caps:
        pf.space_before = Pt(0)
        pf.space_after = Pt(TITLE_SPACE_PT)
    elif chapter_heading:
        pf.space_before = Pt(TITLE_SPACE_PT)
        pf.space_after = Pt(TITLE_SPACE_PT)
    else:
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
    pf.line_spacing = LINE_SPACING
    if toc_entry or (not body and not section_caps and not chapter_heading):
        pf.first_line_indent = Cm(0)
    else:
        pf.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
    if page_break_before:
        pf.page_break_before = True
    if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def classify_line(line: str) -> str:
    s = line.strip()
    if not s:
        return "empty"
    if not is_heading_line(s):
        return "body"
    if is_section_caps_title(s):
        return "section_caps"
    if is_chapter_heading(s):
        return "chapter"
    if s.startswith("Рисунок"):
        return "caption"
    if s.startswith("- ") or s.startswith("— "):
        return "bullet"
    return "body"


@dataclass
class IngestContext:
    has_content: bool = False
    in_toc: bool = False


def ingest_txt(
    doc: Document,
    text: str,
    state: WordBuildState,
    ctx: IngestContext | None = None,
) -> IngestContext:
    if ctx is None:
        ctx = IngestContext()

    for raw in text.splitlines():
        s = raw.strip()
        if not s:
            continue

        if s == TOC_START:
            ctx.in_toc = True
            add_heading(doc, s, level=1, state=state)
            apply_paragraph_format(doc.paragraphs[-1], body=False, section_caps=True)
            ctx.has_content = True
            continue

        if ctx.in_toc:
            if is_section_caps_title(s) and s != TOC_START:
                ctx.in_toc = False
            else:
                add_paragraph(doc, s, state=state)
                apply_paragraph_format(doc.paragraphs[-1], body=False, toc_entry=True)
                ctx.has_content = True
                continue

        kind = classify_line(raw)
        page_break = kind == "section_caps" and ctx.has_content

        if kind == "section_caps":
            add_heading(doc, s, level=1, state=state)
            apply_paragraph_format(
                doc.paragraphs[-1],
                body=False,
                page_break_before=page_break,
                section_caps=True,
            )
        elif kind == "chapter":
            add_heading(doc, s, level=2, state=state)
            apply_paragraph_format(doc.paragraphs[-1], body=False, chapter_heading=True)
        elif kind == "caption":
            add_paragraph(doc, s, state=state)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            apply_paragraph_format(doc.paragraphs[-1], body=False)
        elif kind == "bullet":
            add_paragraph(doc, s[2:].strip(), style="List Bullet", state=state)
            apply_paragraph_format(doc.paragraphs[-1], body=False)
        else:
            add_paragraph(doc, s, state=state)
            apply_paragraph_format(doc.paragraphs[-1], body=True)

        ctx.has_content = True

    return ctx


def build_all(base_dir: Path) -> list[Path]:
    output_dir = base_dir / "docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    results: list[Path] = []

    for txt_name, stem in TXT_FILES:
        src = base_dir / txt_name
        if not src.is_file():
            raise FileNotFoundError(src)

        cfg = WordDocumentConfig(
            output_dir=output_dir,
            filename_stem=stem,
            title=stem,
            author="",
            paper="A4",
            orientation="portrait",
        )
        state = WordBuildState()
        doc = create_document(cfg, state)
        apply_page_setup(doc)
        apply_base_styles(doc)
        ingest_txt(doc, src.read_text(encoding="utf-8"), state)
        format_all_runs(doc)
        apply_title_spacing(doc)
        apply_page_border(doc, state)
        out = cfg.output_path()
        save_docx(doc, out)
        results.append(out)

    cfg = WordDocumentConfig(
        output_dir=output_dir,
        filename_stem=COMBINED_STEM,
        title=COMBINED_STEM,
        author="",
        paper="A4",
        orientation="portrait",
    )
    state = WordBuildState()
    doc = create_document(cfg, state)
    apply_page_setup(doc)
    apply_base_styles(doc)
    ctx = IngestContext()
    for txt_name in COMBINED_TXT_ORDER:
        src = base_dir / txt_name
        ctx = ingest_txt(doc, src.read_text(encoding="utf-8"), state, ctx)
    format_all_runs(doc)
    apply_title_spacing(doc)
    apply_page_border(doc, state)
    combined = cfg.output_path()
    save_docx(doc, combined)
    results.append(combined)
    return results
